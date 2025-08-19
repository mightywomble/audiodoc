from flask import Flask, request, jsonify, render_template, send_file
import whisper
import os
import tempfile
from docx import Document
import io
import uuid
from datetime import datetime
import logging
import re

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['REQUEST_TIMEOUT'] = 600  # 10 minute timeout for large files

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load Whisper model - will be loaded lazily to avoid startup issues
model = None

def get_whisper_model():
    """Lazy loading of Whisper model with optimizations for Apple Silicon"""
    global model
    if model is None:
        logger.info("Loading Whisper large-v3 model (this may take a moment on first run)...")
        # Load the large-v3 model with Apple Silicon optimizations
        model = whisper.load_model(
            "large-v3",
            device="cpu",  # Use CPU for Apple Silicon optimization
            download_root=None  # Use default cache location
        )
        logger.info("Whisper large-v3 model loaded successfully")
    return model

# Ensure upload directory exists
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def format_with_whisper_segments(result):
    """
    Format transcribed text using Whisper's segment data and word timestamps for natural paragraph breaks.
    """
    try:
        if not result or 'segments' not in result:
            return result.get('text', '').strip() if result else ''
        
        segments = result['segments']
        if not segments:
            return result.get('text', '').strip()
        
        paragraphs = []
        current_paragraph = []
        
        for i, segment in enumerate(segments):
            segment_text = segment['text'].strip()
            
            if not segment_text:
                continue
            
            # Analyze segment for paragraph breaks
            should_start_new_paragraph = False
            
            # Start new paragraph if:
            # 1. It's the first segment
            # 2. There's a significant pause (gap in timestamps)
            # 3. Current paragraph is getting long
            # 4. Segment starts with transitional words
            
            if i == 0:
                should_start_new_paragraph = True
            else:
                prev_segment = segments[i-1]
                
                # Check for pause between segments (longer than 2 seconds indicates topic change)
                pause_duration = segment['start'] - prev_segment['end']
                if pause_duration > 2.0:
                    should_start_new_paragraph = True
                
                # Check if current paragraph is getting too long (more than 3 segments)
                if len(current_paragraph) >= 3:
                    should_start_new_paragraph = True
                
                # Check for transitional phrases at start of segment
                first_words = segment_text.lower().split()[:3]
                transition_words = {
                    'now', 'next', 'then', 'however', 'furthermore', 'moreover', 
                    'additionally', 'meanwhile', 'therefore', 'consequently', 
                    'in conclusion', 'finally', 'first', 'second', 'third',
                    'on the other hand', 'in contrast', 'similarly', 'for example'
                }
                
                for word_combo in [' '.join(first_words[:j]) for j in range(1, len(first_words) + 1)]:
                    if word_combo in transition_words:
                        should_start_new_paragraph = True
                        break
            
            # Start new paragraph if needed
            if should_start_new_paragraph and current_paragraph:
                paragraphs.append(' '.join(current_paragraph))
                current_paragraph = []
            
            # Clean up segment text
            cleaned_text = segment_text.strip()
            if cleaned_text:
                # Ensure proper capitalization
                if not current_paragraph:  # First sentence of paragraph
                    cleaned_text = cleaned_text[0].upper() + cleaned_text[1:] if len(cleaned_text) > 1 else cleaned_text.upper()
                
                current_paragraph.append(cleaned_text)
        
        # Add the last paragraph
        if current_paragraph:
            paragraphs.append(' '.join(current_paragraph))
        
        # Join paragraphs with double line breaks
        formatted_text = '\n\n'.join(paragraphs)
        
        # Final cleanup
        # Clean up spaces within paragraphs, but preserve paragraph breaks
        lines = formatted_text.split('\n')
        cleaned_lines = []
        for line in lines:
            if line.strip():  # Non-empty line
                cleaned_lines.append(re.sub(r' +', ' ', line.strip()))  # Multiple spaces to single space
            else:  # Empty line (paragraph separator)
                cleaned_lines.append('')
        formatted_text = '\n'.join(cleaned_lines)
        formatted_text = re.sub(r'\n{3,}', '\n\n', formatted_text)  # Remove excessive line breaks
        
        logger.info(f"Formatted text using Whisper segments into {len(paragraphs)} paragraphs")
        return formatted_text
        
    except Exception as e:
        logger.error(f"Error in Whisper-based formatting: {e}")
        # Fallback to basic text formatting
        return result.get('text', '').strip() if result else ''

def format_transcribed_text(text):
    """
    Format transcribed text into proper paragraphs and sentences for better readability.
    """
    if not text or not text.strip():
        return text
    
    # Clean up the text first
    formatted_text = text.strip()
    
    # Fix common transcription issues
    formatted_text = re.sub(r'\s+', ' ', formatted_text)  # Multiple spaces to single space
    formatted_text = re.sub(r'([.!?])([A-Z])', r'\1 \2', formatted_text)  # Add space after sentence endings
    
    # Add periods to sentences that end without punctuation
    # Look for word followed by capital letter (likely new sentence)
    formatted_text = re.sub(r'([a-z])\s+([A-Z][a-z])', r'\1. \2', formatted_text)
    
    # Split into sentences
    sentences = re.split(r'[.!?]+', formatted_text)
    sentences = [s.strip() for s in sentences if s.strip()]
    
    if not sentences:
        return formatted_text
    
    # Group sentences into paragraphs based on natural breaks
    paragraphs = []
    current_paragraph = []
    
    # Topic change indicators (common words that suggest new topics)
    topic_change_words = {
        'now', 'next', 'then', 'first', 'second', 'third', 'finally', 'meanwhile', 
        'however', 'furthermore', 'moreover', 'additionally', 'in conclusion',
        'on the other hand', 'in contrast', 'similarly', 'likewise', 'therefore',
        'consequently', 'as a result', 'for example', 'for instance', 'in fact',
        'indeed', 'actually', 'specifically', 'particularly', 'especially'
    }
    
    for i, sentence in enumerate(sentences):
        if not sentence.strip():
            continue
            
        # Capitalize first letter of sentence
        sentence = sentence.strip()
        if sentence:
            sentence = sentence[0].upper() + sentence[1:] if len(sentence) > 1 else sentence.upper()
        
        # Check if this sentence should start a new paragraph
        start_new_paragraph = False
        
        # Start new paragraph if:
        # 1. It's the first sentence
        # 2. Current paragraph already has 4+ sentences (avoid very long paragraphs)
        # 3. Sentence starts with a topic change word
        # 4. There's a significant pause indicated by length difference
        
        if i == 0:
            start_new_paragraph = True
        elif len(current_paragraph) >= 4:
            start_new_paragraph = True
        else:
            # Check if sentence starts with topic change words
            first_words = sentence.lower().split()[:3]
            for word_combo in [' '.join(first_words[:i]) for i in range(1, len(first_words) + 1)]:
                if word_combo in topic_change_words:
                    start_new_paragraph = True
                    break
            
            # Check for significant length difference (might indicate pause/topic change)
            if current_paragraph:
                avg_length = sum(len(s) for s in current_paragraph) / len(current_paragraph)
                if len(sentence) > avg_length * 1.5 or len(sentence) < avg_length * 0.5:
                    # Only start new paragraph if current one has at least 2 sentences
                    if len(current_paragraph) >= 2:
                        start_new_paragraph = True
        
        if start_new_paragraph and current_paragraph:
            paragraphs.append(current_paragraph)
            current_paragraph = []
        
        current_paragraph.append(sentence)
    
    # Don't forget the last paragraph
    if current_paragraph:
        paragraphs.append(current_paragraph)
    
    # Join sentences in paragraphs and paragraphs with double line breaks
    formatted_paragraphs = []
    for paragraph in paragraphs:
        # Join sentences with proper punctuation
        paragraph_text = []
        for sentence in paragraph:
            sentence = sentence.strip()
            if sentence:
                # Add period if sentence doesn't end with punctuation
                if not sentence[-1] in '.!?':
                    sentence += '.'
                paragraph_text.append(sentence)
        
        if paragraph_text:
            formatted_paragraphs.append(' '.join(paragraph_text))
    
    # Join paragraphs with double line breaks
    result = '\n\n'.join(formatted_paragraphs)
    
    # Final cleanup
    result = re.sub(r'\n{3,}', '\n\n', result)  # Remove excessive line breaks
    result = re.sub(r'\s+', ' ', result)  # Clean up any remaining multiple spaces
    result = result.replace('\n ', '\n')  # Remove spaces at beginning of lines
    
    logger.info(f"Formatted text into {len(formatted_paragraphs)} paragraphs")
    return result

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        file = request.files['audio']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        filename = f"{file_id}_{file.filename}"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        
        # Save uploaded file
        file.save(filepath)
        logger.info(f"File saved: {filepath}")
        
        # Convert audio to text
        text = convert_audio_to_text(filepath)
        
        # Clean up uploaded file
        os.remove(filepath)
        
        if text:
            return jsonify({
                'success': True,
                'text': text,
                'file_id': file_id,
                'timestamp': datetime.now().isoformat()
            })
        else:
            return jsonify({'error': 'Could not extract text from audio file'}), 400
            
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}")
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500

def convert_audio_to_text(filepath):
    import psutil
    import gc
    import torch
    
    try:
        logger.info(f"Starting transcription for: {filepath}")
        
        # Log system memory info
        memory = psutil.virtual_memory()
        logger.info(f"System memory: {memory.available / (1024**3):.1f}GB available, {memory.percent}% used")
        
        # Validate file exists and has content
        if not os.path.exists(filepath):
            logger.error(f"File does not exist: {filepath}")
            return None
            
        file_size = os.path.getsize(filepath)
        if file_size == 0:
            logger.error(f"File is empty: {filepath}")
            return None
            
        logger.info(f"File size: {file_size} bytes ({file_size / (1024**2):.1f}MB)")
        
        # Load audio and check duration to avoid problematic files
        try:
            # Load audio first to check if it's valid
            audio = whisper.load_audio(filepath)
            duration = len(audio) / whisper.audio.SAMPLE_RATE
            logger.info(f"Audio duration: {duration:.2f} seconds")
            
            # Skip very short audio files that might cause tensor issues
            if duration < 0.5:  # Less than 0.5 seconds
                logger.error("Audio file too short (minimum 0.5 seconds)")
                return None
                
            # Skip very long audio files to avoid memory issues
            if duration > 1800:  # More than 30 minutes
                logger.error("Audio file too long (max 30 minutes)")
                return None
                
        except Exception as audio_load_error:
            logger.error(f"Failed to load audio: {audio_load_error}")
            return None
        
        # Clear any existing model and force garbage collection
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        gc.collect()
        
        # Try multiple strategies for transcription
        strategies = [
            {
                'name': 'Full Quality (large-v3 with word timestamps)',
                'model': 'large-v3',
                'params': {
                    'fp16': False,
                    'language': "en",
                    'task': "transcribe",
                    'verbose': False,
                    'temperature': 0.0,
                    'word_timestamps': True,
                    'condition_on_previous_text': True,
                    'no_speech_threshold': 0.6,
                    'logprob_threshold': -1.0,
                    'compression_ratio_threshold': 2.4,
                }
            },
            {
                'name': 'Simplified (large-v3 basic)',
                'model': 'large-v3',
                'params': {
                    'fp16': False,
                    'language': "en",
                    'task': "transcribe",
                    'verbose': False,
                    'word_timestamps': False,
                    'condition_on_previous_text': False
                }
            },
            {
                'name': 'Medium Model Fallback',
                'model': 'medium.en',
                'params': {
                    'fp16': False,
                    'language': "en",
                    'task': "transcribe",
                    'verbose': False,
                    'word_timestamps': False,
                    'condition_on_previous_text': False
                }
            },
            {
                'name': 'Base Model Fallback',
                'model': 'base.en',
                'params': {
                    'fp16': False,
                    'language': "en",
                    'task': "transcribe",
                    'verbose': False,
                    'word_timestamps': False,
                    'condition_on_previous_text': False
                }
            }
        ]
        
        for i, strategy in enumerate(strategies):
            try:
                logger.info(f"Attempting strategy {i+1}/4: {strategy['name']}")
                
                # Load the appropriate model
                if strategy['model'] != 'large-v3' or model is None:
                    logger.info(f"Loading {strategy['model']} model...")
                    current_model = whisper.load_model(
                        strategy['model'],
                        device="cpu",
                        download_root=None
                    )
                else:
                    current_model = get_whisper_model()
                
                # Clear cache before transcription
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                gc.collect()
                
                # Attempt transcription
                result = current_model.transcribe(filepath, **strategy['params'])
                
                text = result["text"].strip()
                logger.info(f"Strategy {i+1} succeeded! Text length: {len(text)}")
                logger.info(f"Number of segments: {len(result.get('segments', []))}")
                
                if not text:
                    logger.warning(f"Strategy {i+1} returned empty text, trying next strategy...")
                    continue
                
                # Format the text based on available data
                if result.get('segments') and strategy['params'].get('word_timestamps', False):
                    formatted_text = format_with_whisper_segments(result)
                    logger.info(f"Used Whisper segment formatting")
                else:
                    formatted_text = format_transcribed_text(text)
                    logger.info(f"Used basic text formatting")
                
                return formatted_text
                
            except Exception as strategy_error:
                logger.error(f"Strategy {i+1} failed: {strategy_error}")
                
                # Clean up memory after failed attempt
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                gc.collect()
                
                # If it's a tensor shape error, try the next strategy immediately
                error_str = str(strategy_error).lower()
                if any(keyword in error_str for keyword in ['tensor', 'reshape', 'linear', 'cuda', 'memory']):
                    logger.info(f"Memory/tensor error detected, trying next strategy...")
                    continue
                
                # For other errors, also continue to next strategy
                continue
        
        # If all strategies failed
        logger.error("All transcription strategies failed")
        return None
                
    except Exception as e:
        logger.error(f"General error in speech recognition: {e}")
        return None

@app.route('/download_word', methods=['POST'])
def download_word():
    try:
        data = request.get_json()
        text = data.get('text', '')
        filename = data.get('filename', 'transcription')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Create Word document with proper paragraph formatting
        doc = Document()
        doc.add_heading('Audio Transcription', 0)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        
        # Split text into paragraphs and add each as separate paragraph in Word
        paragraphs = text.split('\n\n')
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())
        
        # Save to memory
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name=f'{filename}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        return jsonify({'error': f'Error creating Word document: {str(e)}'}), 500

@app.route('/format_text', methods=['POST'])
def format_text():
    """Endpoint to format raw text into paragraphs"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        formatted_text = format_transcribed_text(text)
        
        return jsonify({
            'success': True,
            'formatted_text': formatted_text
        })
        
    except Exception as e:
        logger.error(f"Error formatting text: {str(e)}")
        return jsonify({'error': f'Error formatting text: {str(e)}'}), 500

@app.route('/health')
def health_check():
    import psutil
    memory = psutil.virtual_memory()
    return jsonify({
        'status': 'healthy',
        'memory_available_gb': round(memory.available / (1024**3), 1),
        'memory_percent_used': memory.percent,
        'whisper_model_loaded': model is not None,
        'timestamp': datetime.now().isoformat()
    })

@app.route('/test')
def test_endpoint():
    """Test endpoint to verify server is working"""
    return jsonify({
        'message': 'Server is running!',
        'timestamp': datetime.now().isoformat(),
        'endpoints': [
            'GET /',
            'POST /upload',
            'POST /download_word',
            'POST /format_text',
            'GET /health',
            'GET /test'
        ]
    })

if __name__ == '__main__':
    import sys
    
    # Check if we should use production server
    if '--production' in sys.argv:
        # Use Gunicorn for production with longer timeout
        print("Starting production server with Gunicorn...")
        print("Run: gunicorn --bind 0.0.0.0:5001 --timeout 600 --workers 1 app:app")
        sys.exit(0)
    else:
        # Development server with threading
        print("Starting development server...")
        print("Note: For longer processing times, use: python app.py --production")
        app.run(
            debug=True,
            host='0.0.0.0',
            port=5001,
            threaded=True  # Enable threading for concurrent requests
        )
